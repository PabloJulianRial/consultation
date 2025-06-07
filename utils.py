# utils.py

import os
import json
import datetime
import sounddevice as sd
import numpy as np
import whisper
from pydub import AudioSegment
from docx import Document

# 1) Load Whisper model once at startup.
#    Choose "small.en" for a reasonable speed/accuracy tradeoff on CPU.
#    If it’s too slow, change to "tiny.en" (faster, less accurate).
WHISPER_MODEL = whisper.load_model("tiny.en")



def load_all_question_sets():
    """
    Reads data/questions/question_sets.json (or any .json in that folder)
    and returns a dict mapping category → list of {id, text}.
    """
    folder = os.path.join(os.path.dirname(__file__), "data", "questions")
    all_files = [f for f in os.listdir(folder) if f.lower().endswith(".json")]
    question_sets = {}
    for fname in all_files:
        path = os.path.join(folder, fname)
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            # If your JSON has a single top-level object of categories, merge it
            if isinstance(data, dict) and all(isinstance(v, list) for v in data.values()):
                # Assume the JSON structure is exactly like question_sets.json above
                for category, qlist in data.items():
                    question_sets[category] = qlist
            else:
                # If you saved each category in its own file, expect { "diagnostic_category": "...", "questions": […] }
                category = data.get("diagnostic_category")
                questions = data.get("questions", [])
                if category:
                    question_sets[category] = questions
    return question_sets


def record_audio_for_question(session_id, question_id):
    """
    Records from the default microphone until the user presses Enter to stop.
    Saves a WAV under data/audio and returns its full path.
    """
    print(f"\n--- Recording for Session {session_id}, Question {question_id} ---")
    print("Press Enter to start recording. Press Enter again to stop.")
    input("Ready? Hit Enter to begin…")
    fs = 16000  # 16 kHz sample rate (Whisper expects ≥16 kHz)
    audio_chunks = []

    def callback(indata, frames, time, status):
        if status:
            print(status, flush=True)
        audio_chunks.append(indata.copy())

    # Start recording until Enter is pressed again
    with sd.InputStream(samplerate=fs, channels=1, callback=callback):
        print("Recording… Press Enter to stop.")
        input()
    audio_data = np.concatenate(audio_chunks, axis=0)

    # Convert NumPy float32 array (−1.0 to 1.0) to int16 PCM and save via pydub
    audio_int16 = (audio_data * 32767).astype(np.int16)
    wav_folder = os.path.join(os.path.dirname(__file__), "data", "audio")
    os.makedirs(wav_folder, exist_ok=True)

    # Sanitize session_id so the filename has no ':' or '-' (invalid on Windows)
    safe_session = session_id.replace(":", "").replace("-", "")
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    wav_filename = f"session_{safe_session}_q_{question_id}_{timestamp}.wav"
    wav_path = os.path.join(wav_folder, wav_filename)

    seg = AudioSegment(
        audio_int16.tobytes(),
        frame_rate=fs,
        sample_width=2,  # 2 bytes for int16
        channels=1
    )
    seg.export(wav_path, format="wav")
    print(f"Saved recording to: {wav_path}")
    return wav_path



def transcribe_audio_file(wav_path):
    """
    Runs Whisper on the given WAV file and returns the transcript string.
    """
    print(f"Transcribing {wav_path} … (this may take a few seconds)")
    result = WHISPER_MODEL.transcribe(wav_path, language="en")
    transcript = result["text"].strip()
    print(f"Transcript:\n{transcript}\n")
    return transcript


def generate_docx_report(session_data):
    """
    Given an in‐memory session_data dict (see gui.py), writes a DOCX report under 'reports/'.
    Returns the path to that .docx.
    """
    reports_folder = os.path.join(os.path.dirname(__file__), "reports")
    os.makedirs(reports_folder, exist_ok=True)

    # Use patient name + start timestamp to build a filename
    safe_name = session_data["patient_name"].replace(" ", "_")
    timestamp = session_data["started_at"].replace(":", "").replace("-", "")
    filename = f"{safe_name}_{timestamp}.docx"
    full_path = os.path.join(reports_folder, filename)

    # If you provided a polished template under templates/report_template.docx, load it
    template_path = os.path.join(os.path.dirname(__file__), "templates", "report_template.docx")
    if os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()

    # Header info
    doc.add_heading("Consultation Report", level=0)
    doc.add_paragraph(f"Patient Name: {session_data['patient_name']}")
    doc.add_paragraph(f"Date of Birth: {session_data['patient_dob']}")
    doc.add_paragraph(f"Diagnostic Category: {session_data['diagnostic_category']}")
    doc.add_paragraph(f"Session Started: {session_data['started_at']}")
    doc.add_paragraph("")  # blank line

    # Questions & Answers section
    doc.add_heading("Questions & Answers", level=1)
    for item in session_data["questions"]:
        qtext = item["question_text"]
        transcript = item.get("transcript", "") or "(no recording/transcript)"
        typed = item.get("typed_answer", "") or "(no typed notes)"

        doc.add_heading(f"Q: {qtext}", level=2)
        doc.add_paragraph("Transcript:")
        doc.add_paragraph(transcript, style="Intense Quote")
        doc.add_paragraph("Typed Notes/Corrections:")
        doc.add_paragraph(typed)
        doc.add_paragraph("")  # blank line

    # General Notes
    doc.add_heading("General Notes", level=1)
    if session_data["general_notes"]:
        for note in session_data["general_notes"]:
            doc.add_paragraph(note)
    else:
        doc.add_paragraph("(none)")

    # Save the file
    doc.save(full_path)
    print(f"Report saved to: {full_path}")
    return full_path
